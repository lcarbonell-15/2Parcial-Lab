from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Estilos globales ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def parrafo(texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(4)
    return p

def formula(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(12)
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def item(texto):
    p = doc.add_paragraph(texto, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def nota(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def tabla(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RESUMEN TEÓRICO')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2do Parcial - Laboratorio de Métodos')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Integrales + Análisis de Inversiones + Pandas')
run.font.size = Pt(14)
run.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════
titulo('ÍNDICE', 1)
items = [
    '1. Integrales Indefinidas',
    '2. Integrales Definidas',
    '3. Aplicaciones: Costos, Ingresos y Beneficio',
    '4. Aplicaciones: Excedentes y Pérdida de Peso Muerto',
    '5. Valor del Dinero en el Tiempo',
    '6. Métricas de Inversión',
    '7. Costo de Capital (CAPM)',
    '8. Fórmulas de Anualidades',
    '9. Pandas - Manipulación de Datos',
    '10. Fórmulario Resumen',
]
for it in items:
    item(it)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. INTEGRALES INDEFINIDAS
# ═══════════════════════════════════════════
titulo('1. Integrales Indefinidas')

parrafo('La integral indefinida es la operación inversa a la derivada. Si F\'(x) = f(x), entonces:')

formula('∫ f(x) dx = F(x) + C')

parrafo('Donde C es la constante de integración, que se determina utilizando una condición inicial (por ejemplo, F(0) = valor conocido).')

titulo('Propiedades fundamentales', 2)
item('Derivar e integrar son operaciones inversas.')
item('La integral de una constante k es: ∫ k dx = kx + C')
item('La integral de x^n es: ∫ x^n dx = x^(n+1)/(n+1) + C  (n ≠ -1)')
item('La integral de e^x es: ∫ e^x dx = e^x + C')
item('La integral de 1/x es: ∫ (1/x) dx = ln|x| + C')

titulo('Procedimiento para resolver', 2)
parrafo('1) Integrar la función para obtener F(x) + C')
parrafo('2) Usar la condición inicial para despejar C')
parrafo('3) Reemplazar C en la expresión general')

nota('Ejemplo: Si IMg = 2x y F(0) = 5, entonces ∫ 2x dx = x² + C. Con F(0) = 5: 0 + C = 5 → C = 5. Resultado: F(x) = x² + 5')

doc.add_page_break()

# ═══════════════════════════════════════════
# 2. INTEGRALES DEFINIDAS
# ═══════════════════════════════════════════
titulo('2. Integrales Definidas')

parrafo('La integral definida calcula el área bajo la curva entre dos puntos a y b. Se resuelve mediante la Regla de Barrow (Primer Teorema Fundamental del Cálculo):')

formula('∫[a,b] f(x) dx = F(b) - F(a)')

titulo('Interpretación geométrica', 2)
item('Representa el área neta bajo la curva f(x) entre x = a y x = b.')
item('Si f(x) > 0, el área es positiva (por encima del eje x).')
item('Si f(x) < 0, el área es negativa (por debajo del eje x).')

titulo('Aproximación de Riemann', 2)
parrafo('La integral definida se puede aproximar como suma de rectángulos:')

formula('∫[a,b] f(x) dx ≈ Σ f(xᵢ) · Δx')

parrafo('Donde Δx = (b - a)/n y n es la cantidad de rectángulos. A mayor n, mejor la aproximación.')

nota('Ejemplo: ∫[0,3] x² dx = 9. Con Riemann (n=5): 8.91 (muy cercano al valor exacto)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. APLICACIONES: COSTOS, INGRESOS, BENEFICIO
# ═══════════════════════════════════════════
titulo('3. Aplicaciones: Costos, Ingresos y Beneficio')

titulo('3.1 Costo Marginal → Costo Total', 2)

formula('C(q) = ∫ CMg(q) dq + F')

parrafo('Donde F es el costo fijo (costo cuando q = 0). Para hallar F, se usa la condición C(0) = F.')

titulo('Costo Medio', 3)
formula('C_med(q) = C(q) / q')
parrafo('El costo medio se minimiza cuando dC_med/dq = 0.')

titulo('3.2 Ingreso Marginal → Ingreso Total', 2)

formula('I(q) = ∫ IMg(q) dq')

parrafo('No se agrega constante porque I(0) = 0 (si no se vende, no hay ingreso).')

titulo('3.3 Beneficio', 2)

formula('B(q) = I(q) - C(q)')
formula('BMg = IMg - CMg')

parrafo('El punto óptimo de beneficio se encuentra cuando BMg = 0, es decir, IMg = CMg.')

titulo('3.4 Equilibrio de Monopolio', 2)
parrafo('El monopolista maximiza beneficio donde IMg = CMg:')
item('Se resuelve IMg(q) = CMg(q) para hallar q*')
item('Se reemplaza q* en la función de demanda para hallar p*')
item('El beneficio es B* = I(q*) - C(q*)')

titulo('3.5 Duopolio de Cournot', 2)
parrafo('En Cournot, cada empresa elige su cantidad considerando la cantidad del rival:')
item('IMg₁ = a - 2bq₁ - bq₂ (la cantidad del rival reduce la pendiente)')
item('Se resuelven las dos ecuaciones de reacción simultáneamente.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. EXCEDENTES Y DWL
# ═══════════════════════════════════════════
titulo('4. Excedentes y Pérdida de Peso Muerto (DWL)')

titulo('4.1 Excedente del Consumidor (EC)', 2)
parrafo('Es la diferencia entre lo que los consumidores están dispuestos a pagar y lo que realmente pagan:')

formula('EC = ∫[0,q*] [Pd(q) - p*] dq')

parrafo('Donde Pd(q) es la función de demanda inversa y p* es el precio de equilibrio.')

titulo('4.2 Excedente del Productor (EP)', 2)
parrafo('Es la diferencia entre el precio que reciben los productores y su costo mínimo:')

formula('EP = ∫[0,q*] [p* - Ps(q)] dq')

parrafo('Donde Ps(q) es la función de oferta inversa.')

titulo('4.3 Equilibrio de Mercado', 2)
parrafo('Para hallar el equilibrio, se iguala oferta y demanda:')
item('Se despeja Pd(q) de la función de demanda Qd(p)')
item('Se despeja Ps(q) de la función de oferta Qs(p)')
item('Se resuelve Pd(q*) = Ps(q*) para hallar q*')
item('Se reemplaza para hallar p*')

titulo('4.4 Impuestos y DWL', 2)
parrafo('Cuando se impone un impuesto t a los productores, la oferta se desplaza:')
formula('Ps^t(q) = Ps(q) + t')
parrafo('El DWL (Dead Weight Loss) es la pérdida de eficiencia:')

formula('DWL = ∫[qt,qt*] [Ps^t(q) - p*] dq + ∫[qt*,q*] [Pd(q) - p*] dq')

parrafo('Donde qt* es la cantidad con impuesto y q* la cantidad original de equilibrio.')

nota('Ejemplo: Con impuesto t = 10, demanda p = 120 - x², oferta p = 32 + 3x: qt* = -3/2 + √321/2, DWL = 13.98')

doc.add_page_break()

# ═══════════════════════════════════════════
# 5. VALOR DEL DINERO EN EL TIEMPO
# ═══════════════════════════════════════════
titulo('5. Valor del Dinero en el Tiempo')

parrafo('El dinero tiene un valor diferente según cuándo se reciba. La preferencia por la liquidez hace que $1 hoy valga más que $1 mañana.')

titulo('5.1 Equivalencia de Tasas', 2)
parrafo('Convierte una tasa de un período a otro:')

formula('i(t₂) = (1 + i(t₁))^(t₂/t₁) - 1')

titulo('5.2 Proporción (TNA ↔ Efectiva)', 2)
parrafo('Relación entre Tasa Nominal Anual (TNA) y tasa efectiva de un período:')

formula('TNA = i_efectiva · (365/n)')
formula('i_efectiva = TNA · (n/365)')

nota('Ejemplo: TNA 40% a 90 días → efectiva 90 días = 40% × 90/365 = 9.863%')

titulo('5.3 Valor Futuro (Capitalización Compuesta)', 2)
formula('Cn = C₀ · (1 + i)^n')

parrafo('Donde C₀ es el capital inicial, i la tasa por período y n la cantidad de períodos.')

titulo('5.4 Valor Presente', 2)
formula('C₀ = Cn / (1 + i)^n')

parrafo('Descuenta un valor futuro a valor actual. Es la operación inversa del valor futuro.')

titulo('5.5 Los 3 Motivos de la Demanda de Dinero (Keynes)', 2)
item('Transacción: para gastos diarios. Depende del ingreso.')
item('Precaución: para imprevistos. Depende del ingreso.')
item('Especulativo: para oportunidades de inversión. Depende de la tasa de interés.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 6. MÉTRICAS DE INVERSIÓN
# ═══════════════════════════════════════════
titulo('6. Métricas de Inversión')

titulo('6.1 Valor Actual Neto (VAN)', 2)
formula('VAN = Σ [Ft / (1+r)^t] - I₀')
parrafo('Donde Ft son los flujos de caja, r la tasa de descuento e I₀ la inversión inicial.')

item('VAN > 0 → Proyecto RENTABLE')
item('VAN < 0 → Proyecto NO RENTABLE')
item('VAN = 0 → Indiferente')

titulo('6.2 Tasa Interna de Retorno (TIR)', 2)
formula('Σ [Ft / (1+TIR)^t] = 0')

item('TIR > r → Proyecto RENTABLE')
item('TIR < r → Proyecto NO RENTABLE')

nota('La TIR es la tasa que hace que el VAN sea exactamente cero.')

titulo('6.3 Tasa Interna de Retorno Modificada (TIRM)', 2)
formula('TIRM = (VF_positivos / VP_negativos)^(1/n) - 1')

parrafo('La TIRM corrige la TIR asumiendo que:')
item('Los flujos positivos se reinvierten a una tasa de reinversión específica.')
item('Los flujos negativos se descuentan a una tasa de financiamiento.')

parrafo('Por eso la TIRM es más realista que la TIR clásica.')

titulo('6.4 Índice de Rentabilidad (IR)', 2)
formula('IR = VAN / Inversión_inicial')

item('IR > 1 → Proyecto RENTABLE')
item('IR < 1 → Proyecto NO RENTABLE')

titulo('6.5 Periodo de Recuperación Simple (PRS)', 2)
parrafo('Es la cantidad de períodos necesarios para recuperar la inversión inicial sumando flujos brutos (sin descuentar).')

item('Es una medida de LIQUIDEZ, no de rentabilidad.')
item('No considera el valor del dinero en el tiempo.')
item('Mientras menor, mejor.')

titulo('6.6 Periodo de Recuperación Descontado (PRD)', 2)
parrafo('Igual que el PRS pero con flujos descontados:')

item('PRD siempre ≥ PRS (porque los flujos descontados son menores).')
item('Considera el valor del dinero en el tiempo.')

titulo('Resumen comparativo', 2)
tabla(
    ['Métrica', 'Fórmula', 'Criterio', '¿Descuenta?', 'numpy-financial'],
    [
        ['VAN', 'Σ Ft/(1+r)^t - I₀', '> 0', 'Sí', 'nf.npv()'],
        ['TIR', 'Σ Ft/(1+TIR)^t = 0', '> r', 'Sí', 'nf.irr()'],
        ['TIRM', '(VF+/VP-)^(1/n) - 1', '> r', 'Sí', 'nf.mirr()'],
        ['IR', 'VAN / I₀', '> 1', 'Sí', 'Manual'],
        ['PRS', 'Σ flujos brutos', 'Cuántos', 'No', 'Manual'],
        ['PRD', 'Σ flujos descont.', 'Cuántos', 'Sí', 'Manual'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 7. COSTO DE CAPITAL
# ═══════════════════════════════════════════
titulo('7. Costo de Capital (CAPM)')

parrafo('El Costo de Capital (Ke o Ko) se calcula con el modelo CAPM:')

formula('Ko = r_f + β(r_m - r_f)')

parrafo('Con riesgo país:')

formula('Ko = r_f + β(r_m - r_f) + RP')

item('r_f: tasa libre de riesgo (ej: bonos del gobierno)')
item('β: beta del proyecto/empresa (sensibilidad al mercado)')
item('r_m: retorno esperado del mercado')
item('RP: riesgo país')

titulo('Ejemplo de conversión', 2)
parrafo('Para convertir Ko anual a mensual:')
formula('Ko_mensual = (1 + Ko_anual)^(30/365) - 1')

nota('Ejemplo: rf = 4%, β = 1.45, rm = 8.5% → Ko = 4% + 1.45(8.5% - 4%) = 10.525% anual')

doc.add_page_break()

# ═══════════════════════════════════════════
# 8. ANUALIDADES
# ═══════════════════════════════════════════
titulo('8. Fórmulas de Anualidades')

parrafo('Una anualidad es una serie de pagos iguales realizados a intervalos regulares.')

titulo('Valor Presente de una Anualidad', 2)
formula('VP = PMT × [1 - (1+r)^(-n)] / r')

titulo('Valor Futuro de una Anualidad', 2)
formula('VF = PMT × [(1+r)^n - 1] / r')

item('PMT: monto de cada pago periódico')
item('r: tasa de interés por período')
item('n: cantidad de períodos')

doc.add_page_break()

# ═══════════════════════════════════════════
# 9. PANDAS
# ═══════════════════════════════════════════
titulo('9. Pandas - Manipulación de Datos')

titulo('9.1 Funciones de Exploración', 2)
tabla(
    ['Función', 'Uso'],
    [
        ['pd.read_csv("archivo.csv")', 'Cargar archivo CSV'],
        ['df.head()', 'Primeras 5 filas'],
        ['df.info()', 'Info general del DataFrame'],
        ['df.describe()', 'Estadísticas descriptivas'],
        ['df.isnull().sum()', 'Contar valores nulos por columna'],
        ['df.shape', 'Dimensiones (filas, columnas)'],
    ]
)

titulo('9.2 Funciones de Filtrado y Orden', 2)
tabla(
    ['Función', 'Uso'],
    [
        ['df[df["col"] > valor]', 'Filtrar filas por condición'],
        ['df.nlargest(n, "col")', 'Top n filas con mayor valor'],
        ['df.nsmallest(n, "col")', 'Bottom n filas'],
        ['df.sort_values("col")', 'Ordenar por columna'],
        ['df.dropna()', 'Eliminar filas con NaN'],
        ['df.drop_duplicates()', 'Eliminar duplicados'],
    ]
)

titulo('9.3 Funciones de Agrupación', 2)
tabla(
    ['Función', 'Uso'],
    [
        ['df.groupby("col")', 'Agrupar por columna'],
        ['df.groupby("col")["col2"].mean()', 'Promedio por grupo'],
        ['df.value_counts()', 'Contar ocurrencias'],
        ['df["col"].map(dict)', 'Mapear valores con diccionario'],
    ]
)

titulo('9.4 Funciones de Limpieza', 2)
tabla(
    ['Función', 'Uso'],
    [
        ['pd.to_numeric(errors="coerce")', 'Convertir a numérico (NaN en errores)'],
        ['df.replace("texto", NaN)', 'Reemplazar valores'],
        ['df.fillna(valor)', 'Rellenar NaN con un valor'],
    ]
)

titulo('9.5 Exportar Datos', 2)
tabla(
    ['Función', 'Uso'],
    [
        ['df.to_excel("out.xlsx")', 'Exportar a Excel'],
        ['pd.ExcelWriter', 'Múltiples hojas en un Excel'],
        ['df.to_csv("out.csv")', 'Exportar a CSV'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 10. FÓRMULARIO RESUMEN
# ═══════════════════════════════════════════
titulo('10. Fórmulario Resumen')

parrafo('Todas las fórmulas clave para el parcial:')

tabla(
    ['Concepto', 'Fórmula'],
    [
        ['Integral indefinida', '∫ f(x) dx = F(x) + C'],
        ['Integral definida', '∫[a,b] f(x) dx = F(b) - F(a)'],
        ['Riemann', 'Σ f(xᵢ) · Δx'],
        ['Costo Total', 'C(q) = ∫ CMg(q) dq + F'],
        ['Ingreso Total', 'I(q) = ∫ IMg(q) dq'],
        ['Beneficio', 'B(q) = I(q) - C(q)'],
        ['Excedente Consumidor', 'EC = ∫[0,q*] [Pd(q) - p*] dq'],
        ['Excedente Productor', 'EP = ∫[0,q*] [p* - Ps(q)] dq'],
        ['Equivalencia de tasas', 'i(t₂) = (1+i(t₁))^(t₂/t₁) - 1'],
        ['Proporción', 'TNA = i_ef × 365/n'],
        ['Valor Futuro', 'Cn = C₀(1+i)^n'],
        ['Valor Presente', 'C₀ = Cn/(1+i)^n'],
        ['VAN', 'Σ Ft/(1+r)^t - I₀'],
        ['TIR', 'Σ Ft/(1+TIR)^t = 0'],
        ['TIRM', '(VF+/VP-)^(1/n) - 1'],
        ['IR', 'VAN / I₀'],
        ['CAPM', 'Ko = rf + β(rm - rf)'],
        ['CAPM + RP', 'Ko = rf + β(rm - rf) + RP'],
        ['VP Anualidad', 'PMT × [1-(1+r)^(-n)] / r'],
        ['VF Anualidad', 'PMT × [(1+r)^n - 1] / r'],
    ]
)

# ═══════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Resumen_Teorico_Parcial2.docx')
doc.save(output_path)
print(f'Word generado: {output_path}')
